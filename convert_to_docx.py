import docx
import os

def markdown_to_docx(md_path, docx_path):
    doc = docx.Document()
    
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith("1. "):
            doc.add_paragraph(line[3:], style='List Number')
        else:
            doc.add_paragraph(line)

    doc.save(docx_path)
    print(f"Successfully created {docx_path}")

if __name__ == "__main__":
    markdown_to_docx('PROJECT_EXPLANATION.md', 'PROJECT_EXPLANATION.docx')
